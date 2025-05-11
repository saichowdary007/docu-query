#!/usr/bin/env python
import os
import sys
import logging
from pathlib import Path
from langchain_core.documents import Document
import pdfplumber
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import local modules
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from rag_core import embed, retrieve, get_chroma_collection

def load_document(file_path):
    """Load document from file path"""
    logger.info(f"Loading document: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == ".pdf":
        docs = []
        try:
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"PDF has {len(pdf.pages)} pages")
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        doc = Document(
                            page_content=text,
                            metadata={"page": i + 1, "source": os.path.basename(file_path)}
                        )
                        docs.append(doc)
        except Exception as e:
            logger.error(f"Error loading PDF: {e}")
        logger.info(f"Loaded {len(docs)} documents from PDF")
        return docs
        
    elif file_ext == ".docx":
        docs = []
        try:
            doc = docx.Document(file_path)
            all_text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text += para.text + "\n"
                    
            if all_text.strip():
                full_doc = Document(
                    page_content=all_text,
                    metadata={"source": os.path.basename(file_path)}
                )
                docs.append(full_doc)
        except Exception as e:
            logger.error(f"Error loading DOCX: {e}")
        logger.info(f"Loaded {len(docs)} documents from DOCX")
        return docs
        
    elif file_ext in [".txt", ".md"]:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if content.strip():
                doc = Document(
                    page_content=content,
                    metadata={"source": os.path.basename(file_path)}
                )
                return [doc]
        except Exception as e:
            logger.error(f"Error loading text file: {e}")
        return []
    
    else:
        logger.error(f"Unsupported file type: {file_ext}")
        return []

def test_document_indexing(file_path):
    """Test document indexing"""
    # Load document
    docs = load_document(file_path)
    if not docs:
        logger.error(f"Failed to load document: {file_path}")
        return False
        
    logger.info(f"Loaded {len(docs)} documents")
    
    # Embed documents
    chunks_count = embed(file_path, docs)
    logger.info(f"Embedded {chunks_count} chunks")
    
    if chunks_count == 0:
        logger.error("Failed to embed documents")
        return False
    
    # Verify collection has documents
    collection = get_chroma_collection()
    if collection:
        try:
            count = collection._collection.count()
            logger.info(f"Collection has {count} documents")
            if count == 0:
                logger.error("Collection is empty after embedding")
                return False
        except Exception as e:
            logger.error(f"Error checking collection count: {e}")
            return False
    else:
        logger.error("Failed to get collection")
        return False
    
    # Test retrieval
    logger.info("Testing retrieval with a sample query")
    result = retrieve("What is the document about?")
    
    if result["type"] == "error":
        logger.error(f"Retrieval error: {result['message']}")
        return False
    
    if result["type"] == "documents":
        logger.info(f"Retrieved {len(result['data'])} documents")
        for doc in result['data']:
            logger.info(f"Source: {doc.metadata.get('source', 'unknown')}")
            logger.info(f"Content snippet: {doc.page_content[:100]}...")
        return True
    
    return False

def main():
    if len(sys.argv) < 2:
        print("Usage: python test_document_indexing.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)
    
    success = test_document_indexing(file_path)
    if success:
        print("Document indexing test successful!")
    else:
        print("Document indexing test failed!")
        sys.exit(1)

if __name__ == "__main__":
    main() 